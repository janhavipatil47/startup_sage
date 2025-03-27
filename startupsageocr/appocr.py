from flask import Flask, request, render_template
import os
import pytesseract
import cv2
from PIL import Image
from docx import Document  # Import for .docx
import pdfplumber  # Import for PDFs
from google.cloud import language_v1

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def analyze_text(text):
    client = language_v1.LanguageServiceClient()
    document = language_v1.Document(content=text, type_=language_v1.Document.Type.PLAIN_TEXT)
    response = client.analyze_entities(document=document)
    return [(entity.name, entity.type_.name) for entity in response.entities]

@app.route("/", methods=["GET", "POST"])
def upload_file():
    extracted_text = ""  # Default empty text
    if request.method == "POST":
        file = request.files["file"]
        if file:
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
            file.save(filepath)

            # **Handle different file types**
            if file.filename.endswith((".png", ".jpg", ".jpeg")):
                # **Extract text from image**
                img = cv2.imread(filepath)
                img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
                pil_img = Image.fromarray(img)
                extracted_text = pytesseract.image_to_string(pil_img)

            elif file.filename.endswith(".docx"):
                # **Extract text from Word Document**
                doc = Document(filepath)
                extracted_text = "\n".join([para.text for para in doc.paragraphs])

            elif file.filename.endswith(".pdf"):
                # **Extract text from PDF using pdfplumber**
                extracted_text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        extracted_text += page.extract_text() + "\n"

            else:
                return "<h2>Unsupported file type!</h2>"

  
    analyzed_result = analyze_text(extracted_text) if extracted_text else []

    return render_template("index.html", text=analyzed_result)

if __name__ == "__main__":
    app.run(debug=True)
